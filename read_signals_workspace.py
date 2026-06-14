import os
import docx

def read_docx(file_path):
    doc = docx.Document(file_path)
    content = []
    
    # Read paragraphs
    for p in doc.paragraphs:
        if p.text.strip():
            content.append(p.text)
            
    # Read tables
    for table in doc.tables:
        content.append("\n--- Table ---")
        for row in table.rows:
            cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            # De-duplicate adjacent cells due to merged cells
            dedup_cells = []
            for c in cells:
                if not dedup_cells or dedup_cells[-1] != c:
                    dedup_cells.append(c)
            content.append(" | ".join(dedup_cells))
        content.append(" ------------\n")
        
    return "\n".join(content)

if __name__ == "__main__":
    possible_paths = [
        r"[PUB] India_runs_data_and_ai_challenge\[PUB] India_runs_data_and_ai_challenge\India_runs_data_and_ai_challenge\redrob_signals_doc.docx",
        r"[PUB] India_runs_data_and_ai_challenge\India_runs_data_and_ai_challenge\redrob_signals_doc.docx",
        r"redrob_signals_doc.docx"
    ]
    
    path = None
    for p in possible_paths:
        if os.path.exists(p):
            path = p
            break
            
    if path:
        print(f"Reading: {path}\n" + "="*40)
        print(read_docx(path))
    else:
        print("Could not find redrob_signals_doc.docx in any standard paths.")
